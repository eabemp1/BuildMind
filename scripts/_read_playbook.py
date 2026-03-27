from docx import Document
p = r"C:\Users\ASUS TUF\Downloads\BuildMind_Playbook_Updated.docx"
doc = Document(p)
parts = []
for para in doc.paragraphs:
    if para.text.strip():
        parts.append(para.text)
for table in doc.tables:
    for row in table.rows:
        parts.append(" | ".join(cell.text.strip() for cell in row.cells))
text = "\n".join(parts)
print(text[12000:16000])
